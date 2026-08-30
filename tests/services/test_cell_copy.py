"""Patch 08 §4 -- the clipboard payload for one rendered table cell."""

import io

import docx
import pytest
from PyQt6.QtGui import QTextCursor

from cockpit.services.cell_copy import OBJECT_REPLACEMENT_CHARACTER, cell_mime_data
from cockpit.services.notes_renderer import RenderPalette, render_build_notes

from tests.services.test_notes_renderer import (
    PALETTE,
    add_picture,
    only_table,
    png_bytes,
    save,
)


@pytest.fixture(autouse=True)
def _qt_app(qapp):
    return qapp


def render(document, tmp_path, name):
    outcome = render_build_notes(save(document, tmp_path, name), PALETTE)
    assert outcome.is_ok(), outcome.err
    return outcome.ok


def test_a_text_only_cell_yields_both_formats(tmp_path):
    document = docx.Document()
    table = document.add_table(1, 2)
    table.cell(0, 0).text = "Torque to 4 in-lb"
    table.cell(0, 1).text = "OTHER CELL"

    rendered = render(document, tmp_path, "text.docx")
    payload = cell_mime_data(only_table(rendered).cellAt(0, 0), rendered.document)

    assert payload.hasText() and payload.hasHtml()
    assert payload.text().strip() == "Torque to 4 in-lb"
    assert "Torque to 4 in-lb" in payload.html()


def test_only_the_clicked_cell_is_copied(tmp_path):
    document = docx.Document()
    table = document.add_table(1, 2)
    table.cell(0, 0).text = "WANTED"
    table.cell(0, 1).text = "NOT WANTED"

    rendered = render(document, tmp_path, "one.docx")
    payload = cell_mime_data(only_table(rendered).cellAt(0, 0), rendered.document)

    assert "WANTED" in payload.text()
    assert "NOT WANTED" not in payload.text()
    assert "NOT WANTED" not in payload.html()


def test_html_inlines_the_image_as_a_data_uri_and_keeps_no_internal_url(tmp_path):
    document = docx.Document()
    table = document.add_table(1, 1)
    add_picture(table.cell(0, 0), png_bytes(60, 40), display_px=(60, 40))

    rendered = render(document, tmp_path, "image.docx")
    assert rendered.image_names, "fixture did not register an image"
    payload = cell_mime_data(only_table(rendered).cellAt(0, 0), rendered.document)

    html = payload.html()
    assert "data:image/png;base64," in html
    assert "notes-img:" not in html
    assert "notes://" not in html


def test_plain_text_carries_no_object_replacement_character(tmp_path):
    document = docx.Document()
    table = document.add_table(1, 1)
    cell = table.cell(0, 0)
    cell.paragraphs[0].add_run("Serial label:")
    add_picture(cell, png_bytes(30, 20), display_px=(30, 20))

    rendered = render(document, tmp_path, "mixed.docx")
    payload = cell_mime_data(only_table(rendered).cellAt(0, 0), rendered.document)

    assert OBJECT_REPLACEMENT_CHARACTER not in payload.text()
    assert "Serial label:" in payload.text()


def test_a_cell_with_several_images_inlines_every_one(tmp_path):
    document = docx.Document()
    table = document.add_table(1, 1)
    cell = table.cell(0, 0)
    add_picture(cell, png_bytes(20, 20, 0xFFFF0000), display_px=(20, 20))
    add_picture(cell, png_bytes(24, 24, 0xFF00FF00), display_px=(24, 24))

    rendered = render(document, tmp_path, "two-images.docx")
    assert len(rendered.image_names) == 2
    payload = cell_mime_data(only_table(rendered).cellAt(0, 0), rendered.document)

    assert payload.html().count("data:image/png;base64,") == 2
    assert "notes-img:" not in payload.html()


def test_the_payload_is_a_fragment_not_a_whole_document(tmp_path):
    """QTextDocument.toHtml would drag body-level styling along with the cell."""
    document = docx.Document()
    table = document.add_table(2, 2)
    table.cell(0, 0).text = "CELL"
    table.cell(1, 1).text = "ELSEWHERE"

    rendered = render(document, tmp_path, "fragment.docx")
    payload = cell_mime_data(only_table(rendered).cellAt(0, 0), rendered.document)

    assert "ELSEWHERE" not in payload.html()


def test_an_empty_cell_still_yields_a_payload(tmp_path):
    document = docx.Document()
    table = document.add_table(1, 1)

    rendered = render(document, tmp_path, "empty.docx")
    payload = cell_mime_data(only_table(rendered).cellAt(0, 0), rendered.document)

    assert payload.hasHtml()
    assert payload.text().strip() == ""
