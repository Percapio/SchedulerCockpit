"""Renders a build-notes .docx into a QTextDocument (Patch 08 §3).

Pure function of a file and a palette value: touches no repository, no cache,
no widget, and no live theme.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Generic, Literal, TypeVar

import docx
from docx.table import _Cell
from PyQt6.QtCore import QBuffer, QByteArray, QIODevice, QSize, QUrl, Qt
from PyQt6.QtGui import (
    QColor,
    QFont,
    QImage,
    QImageReader,
    QPainter,
    QPen,
    QTextCharFormat,
    QTextCursor,
    QTextDocument,
    QTextFrameFormat,
    QTextImageFormat,
    QTextLength,
    QTextTableCellFormat,
    QTextTableFormat,
)

logger = logging.getLogger(__name__)

T = TypeVar("T")
E = TypeVar("E")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

_DRAWING_NS = {"a": A_NS, "wp": WP_NS}

EMU_PER_PX = 9525          # 914400 EMU per inch / 96 px per inch
TWIPS_PER_PX = 15.0        # 1440 twips per inch / 96 px per inch

DECODABLE_CONTENT_TYPES = frozenset({
    "image/png",
    "image/jpeg",
    "image/jpg",
    "image/gif",
    "image/bmp",
    "image/tiff",
    "image/webp",
})

MAX_DOCUMENT_MEDIA_PX = 16_000_000
MAX_IMAGE_EDGE_PX = 4096


def _w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


@dataclass
class Result(Generic[T, E]):
    ok: T | None = None
    err: E | None = None

    def is_ok(self) -> bool:
        return self.ok is not None


@dataclass(frozen=True)
class RenderPalette:
    """Colours the renderer needs that the document does not supply.

    Snapshotted from Theme at the call site and passed by value; the renderer
    never reads live theme state.
    """
    page_background_rgb: str
    default_text_rgb: str
    placeholder_border_rgb: str
    placeholder_text_rgb: str


AnomalyReason = Literal[
    "UndecodableMedium",
    "DegenerateImage",
    "MergeConflict",
    "MediaBudgetExceeded",
]


@dataclass(frozen=True)
class RenderAnomaly:
    """Something in the source that could not be reproduced."""
    reason: AnomalyReason
    relationship: str | None
    table_index: int
    row: int
    column: int
    detail: str = ""


@dataclass
class RenderedNotes:
    document: QTextDocument
    natural_width: int
    anomalies: list[RenderAnomaly] = field(default_factory=list)
    image_names: tuple[str, ...] = ()


@dataclass(frozen=True)
class NotesRenderFailure:
    reason: Literal["DocumentMissing", "DocumentUnreadable", "NoTablesFound"]
    message: str = ""


# ---------------------------------------------------------------- geometry


def _grid_widths_px(docx_table: Any, column_count: int) -> list[float]:
    """Per-column widths in px from w:tblGrid; 0.0 where the column declares none.

    tblGrid coverage is not all-or-nothing, so a missing or zero w:w is
    reported as 0.0 rather than dropped -- the caller distinguishes the three
    cases of §3.2.
    """
    try:
        grid = docx_table._tbl.tblGrid
    except Exception:
        grid = None          # python-docx raises when w:tblGrid is absent
    widths: list[float] = []
    if grid is not None:
        for grid_col in grid.gridCol_lst:
            declared = grid_col.w
            widths.append(declared.twips / TWIPS_PER_PX if declared else 0.0)

    widths += [0.0] * max(0, column_count - len(widths))
    return widths[:column_count]


def _declared_table_width_px(docx_table: Any) -> float:
    """w:tblW in px when expressed in dxa (twentieths of a point); 0.0 otherwise."""
    table_properties = docx_table._tbl.tblPr
    if table_properties is None:
        return 0.0
    declared = table_properties.find(_w("tblW"))
    if declared is None:
        return 0.0
    if declared.get(_w("type")) != "dxa":
        return 0.0
    try:
        return int(declared.get(_w("w"), 0)) / TWIPS_PER_PX
    except (TypeError, ValueError):
        return 0.0


def resolve_column_constraints(
    docx_table: Any,
    column_count: int
) -> tuple[list[QTextLength], float]:
    """Column width constraints and the table's pixel extent (§3.2).

    post: the second element is 0.0 for a table that fell back to percentage
          widths -- such a table has no pixel extent until it is laid out and
          must not contribute to natural_width
    """
    if column_count <= 0:
        return [], 0.0

    widths = _grid_widths_px(docx_table, column_count)
    specified = [w for w in widths if w > 0]

    if not specified:
        even = 100.0 / column_count
        return (
            [QTextLength(QTextLength.Type.PercentageLength, even)] * column_count,
            0.0,
        )

    if len(specified) < column_count:
        declared_total = _declared_table_width_px(docx_table)
        specified_total = sum(specified)
        unspecified_count = column_count - len(specified)
        if declared_total > specified_total:
            fill = (declared_total - specified_total) / unspecified_count
        else:
            fill = specified_total / len(specified)
        widths = [w if w > 0 else fill for w in widths]

    return (
        [QTextLength(QTextLength.Type.FixedLength, w) for w in widths],
        float(sum(widths)),
    )


# ---------------------------------------------------------------- borders


def _edge_value(borders_element: Any, edge: str) -> str | None:
    if borders_element is None:
        return None
    edge_element = borders_element.find(_w(edge))
    if edge_element is None:
        return None
    return edge_element.get(_w("val")) or "single"


def resolve_cell_borders(cell: Any, docx_table: Any) -> dict[str, bool]:
    """Per-edge border presence, cell over table (§3.3).

    post: an edge explicitly set to nil/none on the cell suppresses and does
          not fall through to the table value; an absent cell edge inherits
    """
    cell_properties = cell._tc.tcPr
    cell_borders = cell_properties.find(_w("tcBorders")) if cell_properties is not None else None
    table_properties = docx_table._tbl.tblPr
    table_borders = table_properties.find(_w("tblBorders")) if table_properties is not None else None

    resolved: dict[str, bool] = {}
    for edge in ("top", "bottom", "left", "right"):
        value = _edge_value(cell_borders, edge)
        if value is None:
            value = _edge_value(table_borders, edge)
        resolved[edge] = bool(value) and value not in ("nil", "none")
    return resolved


# ---------------------------------------------------------------- images


def display_size_px(drawing: Any, reader: QImageReader) -> QSize:
    """Display size of an embedded image, as Word lays it out.

    post: wp:extent (or a:ext) cx/cy converted at 96 DPI; when no extent is
          present, the medium's intrinsic size as reported by QImageReader
    """
    for path in (".//wp:extent", ".//a:ext"):
        extent = drawing.find(path, namespaces=_DRAWING_NS)
        if extent is None:
            continue
        try:
            cx = int(extent.get("cx", 0))
            cy = int(extent.get("cy", 0))
        except (TypeError, ValueError):
            continue
        if cx > 0 and cy > 0:
            return QSize(max(1, round(cx / EMU_PER_PX)), max(1, round(cy / EMU_PER_PX)))

    intrinsic = reader.size()
    if intrinsic.isValid() and intrinsic.width() > 0 and intrinsic.height() > 0:
        return intrinsic
    return QSize()


def _placeholder_image(text: str, palette: RenderPalette) -> QImage:
    """A bordered box naming what could not be rendered (§7)."""
    metrics_font = QFont()
    metrics_font.setPointSize(9)
    width = max(120, 8 * len(text) + 16)
    image = QImage(width, 28, QImage.Format.Format_ARGB32_Premultiplied)
    image.fill(QColor(palette.page_background_rgb))
    painter = QPainter(image)
    try:
        painter.setFont(metrics_font)
        painter.setPen(QPen(QColor(palette.placeholder_border_rgb), 1))
        painter.drawRect(0, 0, image.width() - 1, image.height() - 1)
        painter.setPen(QPen(QColor(palette.placeholder_text_rgb)))
        painter.drawText(
            image.rect().adjusted(6, 0, -6, 0),
            int(Qt.AlignmentFlag.AlignVCenter | Qt.AlignmentFlag.AlignLeft),
            text,
        )
    finally:
        painter.end()
    return image


class _MediaBudget:
    """Total decoded pixels across one document (§3.4)."""

    def __init__(self, ceiling_px: int) -> None:
        self._ceiling_px = ceiling_px
        self._spent_px = 0

    def can_afford(self, size: QSize) -> bool:
        return self._spent_px + size.width() * size.height() <= self._ceiling_px

    def charge(self, image: QImage) -> None:
        self._spent_px += image.width() * image.height()


# ---------------------------------------------------------------- rendering


class _CellPainter:
    """Writes one Word cell's runs into a QTextCursor positioned in a QTextTable."""

    def __init__(
        self,
        document: QTextDocument,
        document_part: Any,
        palette: RenderPalette,
        budget: _MediaBudget,
        anomalies: list[RenderAnomaly],
        image_names: set[str],
    ) -> None:
        self._document = document
        self._document_part = document_part
        self._palette = palette
        self._budget = budget
        self._anomalies = anomalies
        self._image_names = image_names

    def paint(
        self,
        cursor: QTextCursor,
        cell: Any,
        table_index: int,
        row: int,
        column: int,
    ) -> None:
        for paragraph_index, paragraph in enumerate(cell.paragraphs):
            if paragraph_index > 0:
                cursor.insertBlock()
            for run in paragraph.runs:
                char_format = self._char_format(run)
                for child in run._r.iterchildren():
                    tag = child.tag.split("}")[-1]
                    if tag in ("t", "tab", "br", "cr", "noBreakHyphen", "ptab"):
                        text = str(child)
                        if text:
                            cursor.setCharFormat(char_format)
                            cursor.insertText(text)
                    elif tag in ("drawing", "pict", "object"):
                        self._insert_medium(
                            cursor, child, char_format, table_index, row, column
                        )

    def _char_format(self, run: Any) -> QTextCharFormat:
        char_format = QTextCharFormat()
        if run.bold:
            char_format.setFontWeight(QFont.Weight.Bold)
        if run.italic:
            char_format.setFontItalic(True)
        if run.underline:
            char_format.setFontUnderline(True)
        colour = run.font.color
        if colour is not None and colour.rgb is not None:
            char_format.setForeground(QColor(f"#{colour.rgb}"))
        else:
            char_format.setForeground(QColor(self._palette.default_text_rgb))
        if run.font.size is not None:
            char_format.setFontPointSize(run.font.size.pt)
        if run.font.name:
            char_format.setFontFamilies([run.font.name])
        return char_format

    def _relationship_id(self, drawing: Any) -> str | None:
        blip = drawing.find(".//a:blip", namespaces=_DRAWING_NS)
        if blip is not None:
            return blip.get(f"{{{R_NS}}}embed")
        for element in drawing.iter():
            if element.tag.endswith("}imagedata"):
                return element.get(f"{{{R_NS}}}id")
        return None

    def _insert_medium(
        self,
        cursor: QTextCursor,
        drawing: Any,
        char_format: QTextCharFormat,
        table_index: int,
        row: int,
        column: int,
    ) -> None:
        relationship_id = self._relationship_id(drawing)
        if not relationship_id or relationship_id not in self._document_part.related_parts:
            return

        part = self._document_part.related_parts[relationship_id]
        content_type = getattr(part, "content_type", "") or ""

        if content_type.lower() not in DECODABLE_CONTENT_TYPES:
            self._placehold(
                cursor, f"[unsupported image: {content_type or 'unknown'}]",
                RenderAnomaly("UndecodableMedium", relationship_id,
                              table_index, row, column, content_type),
            )
            return

        try:
            blob = part.blob
        except Exception as exc:
            self._placehold(
                cursor, "[unreadable image]",
                RenderAnomaly("UndecodableMedium", relationship_id,
                              table_index, row, column, str(exc)),
            )
            return

        # QBuffer holds a pointer, not a copy: the QByteArray must outlive it.
        payload = QByteArray(blob)
        buffer = QBuffer(payload)
        buffer.open(QIODevice.OpenModeFlag.ReadOnly)
        try:
            reader = QImageReader(buffer)
            size = display_size_px(drawing, reader)
            if size.isEmpty():
                self._placehold(
                    cursor, "[unreadable image]",
                    RenderAnomaly("DegenerateImage", relationship_id,
                                  table_index, row, column, "no usable extent"),
                )
                return

            size = QSize(
                min(size.width(), MAX_IMAGE_EDGE_PX),
                min(size.height(), MAX_IMAGE_EDGE_PX),
            )
            if not self._budget.can_afford(size):
                self._placehold(
                    cursor, "[image omitted: document media budget]",
                    RenderAnomaly("MediaBudgetExceeded", relationship_id,
                                  table_index, row, column),
                )
                return

            reader.setScaledSize(size)
            image = reader.read()
        finally:
            buffer.close()

        if image.isNull() or image.width() <= 0 or image.height() <= 0:
            self._placehold(
                cursor, "[unreadable image]",
                RenderAnomaly("DegenerateImage", relationship_id,
                              table_index, row, column, reader.errorString()),
            )
            return

        name = f"notes-img:{getattr(part, 'sha1', relationship_id)}"
        self._document.addResource(
            QTextDocument.ResourceType.ImageResource, QUrl(name), image
        )
        self._image_names.add(name)
        self._budget.charge(image)

        image_format = QTextImageFormat()
        image_format.setName(name)
        image_format.setWidth(image.width())
        image_format.setHeight(image.height())
        cursor.insertImage(image_format)
        cursor.setCharFormat(char_format)

    def _placehold(
        self,
        cursor: QTextCursor,
        label: str,
        anomaly: RenderAnomaly,
    ) -> None:
        self._anomalies.append(anomaly)
        name = f"notes-placeholder:{label}"
        if name not in self._image_names:
            self._document.addResource(
                QTextDocument.ResourceType.ImageResource,
                QUrl(name),
                _placeholder_image(label, self._palette),
            )
            self._image_names.add(name)
        image_format = QTextImageFormat()
        image_format.setName(name)
        cursor.insertImage(image_format)


def _cell_format(cell: Any, docx_table: Any) -> QTextTableCellFormat:
    cell_format = QTextTableCellFormat()
    solid = QTextFrameFormat.BorderStyle.BorderStyle_Solid
    none = QTextFrameFormat.BorderStyle.BorderStyle_None
    borders = resolve_cell_borders(cell, docx_table)
    cell_format.setTopBorderStyle(solid if borders["top"] else none)
    cell_format.setBottomBorderStyle(solid if borders["bottom"] else none)
    cell_format.setLeftBorderStyle(solid if borders["left"] else none)
    cell_format.setRightBorderStyle(solid if borders["right"] else none)

    cell_properties = cell._tc.tcPr
    if cell_properties is not None:
        shading = cell_properties.find(_w("shd"))
        if shading is not None:
            fill = shading.get(_w("fill"))
            if fill and fill.lower() not in ("auto", "ffffff"):
                cell_format.setBackground(QColor(f"#{fill}"))

        alignment = cell_properties.find(_w("vAlign"))
        value = alignment.get(_w("val")) if alignment is not None else None
    else:
        value = None

    if value == "center":
        cell_format.setVerticalAlignment(QTextCharFormat.VerticalAlignment.AlignMiddle)
    elif value == "bottom":
        cell_format.setVerticalAlignment(QTextCharFormat.VerticalAlignment.AlignBottom)
    else:
        cell_format.setVerticalAlignment(QTextCharFormat.VerticalAlignment.AlignTop)
    return cell_format


def _column_count(rows: list) -> int:
    """Grid width of a table, derived from its cells rather than w:tblGrid.

    Table.columns raises when w:tblGrid is absent, which is one of the three
    cases §3.2 has to handle, so the count comes from the cells themselves.
    """
    widest = 0
    for row_element in rows:
        edge = 0
        for tc in row_element.tc_lst:
            edge = max(edge, tc.grid_offset + max(1, tc.grid_span))
        widest = max(widest, edge)
    return widest


def _apply_merges(
    qtable: Any,
    regions: dict[tuple[int, int], tuple[int, int]],
    table_index: int,
    anomalies: list[RenderAnomaly],
) -> None:
    """Applies collected merge regions after the table is fully populated (§3.2).

    pre:  every region is keyed by its anchor cell in source coordinates
    post: a region overlapping one already applied is dropped and recorded as
          MergeConflict; the cell renders unmerged and the table stays readable
    """
    claimed: set[tuple[int, int]] = set()
    for (row, column) in sorted(regions):
        row_span, column_span = regions[(row, column)]
        if row_span <= 1 and column_span <= 1:
            continue
        cells = [
            (r, c)
            for r in range(row, row + row_span)
            for c in range(column, column + column_span)
        ]
        if any(cell in claimed for cell in cells):
            anomalies.append(
                RenderAnomaly("MergeConflict", None, table_index, row, column,
                              "overlaps an applied merge")
            )
            continue
        claimed.update(cells)
        qtable.mergeCells(row, column, row_span, column_span)


def render_build_notes(
    docx_path: Path,
    palette: RenderPalette,
) -> Result[RenderedNotes, NotesRenderFailure]:
    """Builds a Qt document from the tables of a Word build-note file.

    pre:  docx_path exists and is readable
    post: on success the document contains one QTextTable per source table, in
          document order, with every decodable image registered as a document
          resource under a name this module chose; every irreproducible element
          is both replaced by an inline placeholder and listed in anomalies
    raises: nothing -- every failure is returned as a value
    """
    docx_path = Path(docx_path)
    try:
        if not docx_path.is_file():
            return Result(err=NotesRenderFailure("DocumentMissing", str(docx_path)))
    except OSError as exc:
        return Result(err=NotesRenderFailure("DocumentUnreadable", str(exc)))

    try:
        source = docx.Document(str(docx_path))
    except Exception as exc:
        return Result(err=NotesRenderFailure("DocumentUnreadable", str(exc)))

    try:
        tables = source.tables
    except Exception as exc:
        return Result(err=NotesRenderFailure("DocumentUnreadable", str(exc)))

    if not tables:
        return Result(err=NotesRenderFailure("NoTablesFound"))

    document = QTextDocument()
    document.setUndoRedoEnabled(False)
    cursor = QTextCursor(document)

    anomalies: list[RenderAnomaly] = []
    image_names: set[str] = set()
    budget = _MediaBudget(MAX_DOCUMENT_MEDIA_PX)
    painter = _CellPainter(
        document, source.part, palette, budget, anomalies, image_names
    )

    natural_width = 0.0
    rendered_any = False

    for table_index, docx_table in enumerate(tables):
        try:
            rows = list(docx_table._tbl.tr_lst)
            row_count = len(rows)
            column_count = _column_count(rows)
        except Exception as exc:
            logger.warning("Build-note table %d is unreadable: %s", table_index, exc)
            continue
        if row_count == 0 or column_count == 0:
            continue

        constraints, table_width = resolve_column_constraints(docx_table, column_count)

        table_format = QTextTableFormat()
        table_format.setBorder(0)
        table_format.setCellSpacing(0)
        table_format.setCellPadding(4)
        table_format.setColumnWidthConstraints(constraints)
        if table_width > 0:
            table_format.setWidth(QTextLength(QTextLength.Type.FixedLength, table_width))
            natural_width = max(natural_width, table_width)

        if rendered_any:
            cursor.movePosition(QTextCursor.MoveOperation.End)
            cursor.insertBlock()
        qtable = cursor.insertTable(row_count, column_count, table_format)
        rendered_any = True

        merge_regions: dict[tuple[int, int], tuple[int, int]] = {}
        open_vertical: dict[int, int] = {}

        # The raw w:tc elements, not Table.rows[].cells: the latter resolves a
        # vMerge continue back to its anchor cell, which erases the very state
        # the merge pass needs to read.
        for row_index, row_element in enumerate(rows):
            for tc in row_element.tc_lst:
                try:
                    column_index = tc.grid_offset
                    column_span = max(1, tc.grid_span)
                    vertical = tc.vMerge
                except Exception as exc:
                    logger.warning(
                        "Build-note table %d row %d cell is unreadable: %s",
                        table_index, row_index, exc,
                    )
                    continue
                if column_index >= column_count:
                    continue

                if vertical == "continue":
                    anchor_row = open_vertical.get(column_index)
                    if anchor_row is None:
                        anomalies.append(
                            RenderAnomaly("MergeConflict", None, table_index,
                                          row_index, column_index,
                                          "vMerge continue with no restart")
                        )
                        continue
                    key = (anchor_row, column_index)
                    _, anchor_span = merge_regions.get(key, (1, column_span))
                    merge_regions[key] = (row_index - anchor_row + 1, anchor_span)
                    continue

                if vertical == "restart":
                    open_vertical[column_index] = row_index
                else:
                    open_vertical.pop(column_index, None)
                if column_span > 1 or vertical == "restart":
                    merge_regions[(row_index, column_index)] = (1, column_span)

                target = qtable.cellAt(row_index, column_index)
                if not target.isValid():
                    continue
                cell = _Cell(tc, docx_table)
                target.setFormat(_cell_format(cell, docx_table))
                painter.paint(
                    target.firstCursorPosition(), cell,
                    table_index, row_index, column_index,
                )

        _apply_merges(qtable, merge_regions, table_index, anomalies)
        cursor.movePosition(QTextCursor.MoveOperation.End)

    if not rendered_any:
        return Result(err=NotesRenderFailure("NoTablesFound"))

    return Result(ok=RenderedNotes(
        document=document,
        natural_width=int(natural_width),
        anomalies=anomalies,
        image_names=tuple(sorted(image_names)),
    ))
