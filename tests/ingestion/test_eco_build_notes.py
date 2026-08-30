"""Patch 08 §2 -- the ECO parser reduced to its ingestion roles.

Cell content and image references are no longer extracted; the pane renders the
.docx directly. What survives is the declared part number and the three gates
that reject a malformed document at drop time.
"""

import docx
import pytest

from cockpit.ingestion.errors import MalformedEcoError
from cockpit.ingestion.parsers import eco_build_notes
from cockpit.ingestion.parsers.results import EcoResult

XRAY_HEADER = ["Find#", "PartNum", "Count", "Ref_Des", "Description"]


def save(document, tmp_path, name):
    path = tmp_path / name
    document.save(str(path))
    return path


def fill_row(table, row, values):
    for column, value in enumerate(values):
        table.cell(row, column).text = value


def test_declared_part_number_comes_from_the_filename_prefix(tmp_path):
    document = docx.Document()
    document.add_table(1, 2)
    path = save(document, tmp_path, "B142006 ECO.docx")

    assert eco_build_notes.parse(path).declared_part_number == "B142006"


def test_row_count_excludes_a_recognised_header_row(tmp_path):
    document = docx.Document()
    table = document.add_table(4, 5)
    fill_row(table, 0, XRAY_HEADER)

    result = eco_build_notes.parse(save(document, tmp_path, "B1 ECO.docx"))
    assert result.row_count == 3
    assert result.raw_table_count == 1


def test_row_count_keeps_every_row_when_there_is_no_header(tmp_path):
    document = docx.Document()
    document.add_table(3, 2)

    assert eco_build_notes.parse(save(document, tmp_path, "B1 ECO.docx")).row_count == 3


def test_row_count_sums_across_tables(tmp_path):
    document = docx.Document()
    build = document.add_table(2, 2)
    build.cell(0, 0).text = "#"
    document.add_paragraph("")
    xray = document.add_table(4, 5)
    fill_row(xray, 0, XRAY_HEADER)

    result = eco_build_notes.parse(save(document, tmp_path, "B1 ECO.docx"))
    assert result.row_count == 1 + 3
    assert result.raw_table_count == 2


def test_result_carries_no_cell_content(tmp_path):
    document = docx.Document()
    table = document.add_table(1, 2)
    table.cell(0, 0).text = "Place a serial label"

    result = eco_build_notes.parse(save(document, tmp_path, "B1 ECO.docx"))
    assert isinstance(result, EcoResult)
    assert not hasattr(result, "items")
    assert not hasattr(result, "cells")


# ------------------------------------------------------------------ the gates


def test_unreadable_document_is_rejected(tmp_path):
    path = tmp_path / "B1 ECO.docx"
    path.write_bytes(b"not a docx at all")

    with pytest.raises(MalformedEcoError) as excinfo:
        eco_build_notes.parse(path)
    assert "UNREADABLE_DOCUMENT" in str(excinfo.value)


def test_table_count_drift_is_rejected(tmp_path):
    document = docx.Document()
    for _ in range(4):
        document.add_table(1, 1)
        document.add_paragraph("")

    with pytest.raises(MalformedEcoError) as excinfo:
        eco_build_notes.parse(save(document, tmp_path, "B1 ECO.docx"))
    assert "TABLE_COUNT_DRIFT" in str(excinfo.value)


def test_three_tables_are_accepted(tmp_path):
    document = docx.Document()
    for _ in range(3):
        document.add_table(1, 1)
        document.add_paragraph("")

    assert eco_build_notes.parse(save(document, tmp_path, "B1 ECO.docx")).raw_table_count == 3


def test_xray_header_drift_is_rejected(tmp_path):
    document = docx.Document()
    table = document.add_table(2, 5)
    fill_row(table, 0, ["Find#", "PartNumber", "Count", "Ref_Des", "Description"])

    with pytest.raises(MalformedEcoError) as excinfo:
        eco_build_notes.parse(save(document, tmp_path, "B1 ECO.docx"))
    assert "XRAY_HEADER_DRIFT" in str(excinfo.value)


def test_the_canonical_xray_header_passes(tmp_path):
    document = docx.Document()
    table = document.add_table(2, 5)
    fill_row(table, 0, XRAY_HEADER)

    assert eco_build_notes.parse(save(document, tmp_path, "B1 ECO.docx")).row_count == 1


def test_a_build_table_is_not_held_to_the_xray_header(tmp_path):
    document = docx.Document()
    table = document.add_table(2, 2)
    fill_row(table, 0, ["#", "Instruction"])

    assert eco_build_notes.parse(save(document, tmp_path, "B1 ECO.docx")).row_count == 1


@pytest.mark.parametrize("name", [
    "B140002 Artis (ITAR)/B140002 ECO (ITAR).docx",
    "B142000 Atlas Devices (ITAR)/B142000 ECO (ITAR).docx",
    "B142006 Angel Aerial Systems/B142006 ECO.docx",
])
def test_real_build_notes_pass_every_gate(name):
    import pathlib

    path = pathlib.Path("backend/data") / name
    if not path.exists():
        pytest.skip(f"reference document not present: {path}")

    result = eco_build_notes.parse(path)
    assert result.declared_part_number == path.name.split()[0]
    assert result.row_count > 0
