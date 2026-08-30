"""Patch 08 §3.1, §5, §7, §8 -- the Build Notes pane."""

import docx
import pytest
from PyQt6.QtCore import QPoint, Qt
from PyQt6.QtGui import QTextDocument

from cockpit.ui.widgets.build_notes_pane import (
    NOT_ATTACHED_MESSAGE,
    NO_TABLES_MESSAGE,
    BuildNotesPane,
)


class StubTheme:
    def notes_page_background_rgb(self) -> str: return "#FFFFFF"
    def notes_default_text_rgb(self) -> str: return "#000000"
    def notes_placeholder_border_rgb(self) -> str: return "#CC4444"
    def notes_placeholder_text_rgb(self) -> str: return "#996666"
    def notes_search_highlight_rgb(self) -> str: return "#FFFF00"


class StubView:
    def __init__(self, notes_docx_path=None):
        self.notes_docx_path = notes_docx_path


def build_notes(tmp_path, name, rows=3, text="Torque to 4 in-lb"):
    document = docx.Document()
    table = document.add_table(rows, 2)
    table.cell(0, 0).text = "1."
    table.cell(0, 1).text = text
    for row in range(1, rows):
        table.cell(row, 1).text = f"{text} step {row}"
    path = tmp_path / name
    document.save(str(path))
    return path


@pytest.fixture
def pane(qtbot):
    widget = BuildNotesPane(StubTheme())
    qtbot.addWidget(widget)
    return widget


# ------------------------------------------------------------------ lifecycle


def test_a_fresh_pane_is_empty(pane):
    assert pane._state == "Empty"
    assert pane._document is None


def test_load_while_hidden_does_not_render(pane, tmp_path):
    pane.load(StubView(build_notes(tmp_path, "a.docx")))
    assert pane._state == "Empty"
    assert pane._document is None


def test_rendering_happens_on_first_show(pane, qtbot, tmp_path):
    pane.load(StubView(build_notes(tmp_path, "a.docx")))
    pane.show()
    qtbot.waitExposed(pane)
    assert pane._state == "Rendered"
    assert "Torque to 4 in-lb" in pane._document.toPlainText()


def test_a_second_switch_to_the_page_does_not_re_render(pane, qtbot, tmp_path):
    pane.load(StubView(build_notes(tmp_path, "a.docx")))
    pane.show()
    qtbot.waitExposed(pane)
    first = pane._document

    pane.hide()
    pane.show()
    qtbot.waitExposed(pane)
    assert pane._document is first


def test_load_while_showing_re_renders_to_the_new_audit(pane, qtbot, tmp_path):
    pane.load(StubView(build_notes(tmp_path, "a.docx", text="AUDIT A")))
    pane.show()
    qtbot.waitExposed(pane)
    assert "AUDIT A" in pane._document.toPlainText()

    pane.load(StubView(build_notes(tmp_path, "b.docx", text="AUDIT B")))
    assert pane._state == "Rendered"
    assert "AUDIT B" in pane._document.toPlainText()
    assert "AUDIT A" not in pane._document.toPlainText()


def test_load_while_the_pdf_page_is_showing_defers_until_the_switch(pane, qtbot, tmp_path):
    pane.load(StubView(build_notes(tmp_path, "a.docx", text="AUDIT A")))
    pane.show()
    qtbot.waitExposed(pane)

    pane.hide()
    pane.load(StubView(build_notes(tmp_path, "b.docx", text="AUDIT B")))
    assert pane._state == "Empty"

    pane.show()
    qtbot.waitExposed(pane)
    assert pane._state == "Rendered"
    assert "AUDIT B" in pane._document.toPlainText()


def test_unload_gives_the_editor_a_fresh_document_and_clears_the_old(pane, qtbot, tmp_path):
    pane.load(StubView(build_notes(tmp_path, "a.docx")))
    pane.show()
    qtbot.waitExposed(pane)
    stale = pane._document

    pane.unload()
    assert pane._state == "Empty"
    assert pane._document is None
    assert pane._editor.document() is not stale
    assert stale.toPlainText() == ""


def test_unload_is_idempotent(pane):
    pane.unload()
    pane.unload()
    assert pane._state == "Empty"


def test_invalidate_returns_the_pane_to_empty_when_hidden(pane, tmp_path):
    pane.load(StubView(build_notes(tmp_path, "a.docx")))
    pane.invalidate()
    assert pane._state == "Empty"


# ------------------------------------------------------------------ failures


def test_no_notes_path_shows_not_attached_without_calling_the_renderer(
    pane, qtbot, monkeypatch
):
    import cockpit.ui.widgets.build_notes_pane as module

    def explode(*args, **kwargs):
        raise AssertionError("renderer must not be called without a path")

    monkeypatch.setattr(module, "render_build_notes", explode)

    pane.load(StubView(None))
    pane.show()
    qtbot.waitExposed(pane)
    assert pane._stack.currentWidget() is pane._message
    assert pane._message.text() == NOT_ATTACHED_MESSAGE


def test_a_missing_document_names_the_path_and_the_remedy(pane, qtbot, tmp_path):
    pane.load(StubView(tmp_path / "gone.docx"))
    pane.show()
    qtbot.waitExposed(pane)
    assert pane._stack.currentWidget() is pane._message
    assert "gone.docx" in pane._message.text()
    assert "re-ingest" in pane._message.text().lower()


def test_an_unreadable_document_shows_the_underlying_reason(pane, qtbot, tmp_path):
    corrupt = tmp_path / "corrupt.docx"
    corrupt.write_bytes(b"PK\x03\x04 nonsense")
    pane.load(StubView(corrupt))
    pane.show()
    qtbot.waitExposed(pane)
    assert pane._stack.currentWidget() is pane._message
    assert "could not be read" in pane._message.text()


def test_a_document_with_no_tables_says_so(pane, qtbot, tmp_path):
    document = docx.Document()
    document.add_paragraph("Prose only.")
    path = tmp_path / "prose.docx"
    document.save(str(path))

    pane.load(StubView(path))
    pane.show()
    qtbot.waitExposed(pane)
    assert pane._message.text() == NO_TABLES_MESSAGE


def test_a_render_failure_never_raises_into_the_caller(pane, tmp_path):
    pane.show()
    pane.load(StubView(tmp_path / "absent.docx"))  # must not raise


# ------------------------------------------------------------------ search


def test_highlight_matches_returns_zero_on_an_empty_pane(pane):
    assert pane.highlight_matches("anything") == 0


def test_highlight_matches_counts_and_marks_every_occurrence(pane, qtbot, tmp_path):
    pane.load(StubView(build_notes(tmp_path, "a.docx", rows=4, text="TORQUE")))
    pane.show()
    qtbot.waitExposed(pane)

    count = pane.highlight_matches("TORQUE")
    assert count == 4
    assert len(pane._editor.extraSelections()) == 4


def test_an_empty_query_clears_the_selections(pane, qtbot, tmp_path):
    pane.load(StubView(build_notes(tmp_path, "a.docx", text="TORQUE")))
    pane.show()
    qtbot.waitExposed(pane)
    pane.highlight_matches("TORQUE")

    assert pane.highlight_matches("") == 0
    assert pane._editor.extraSelections() == []


def test_a_query_that_matches_nothing_returns_zero_and_keeps_the_document(
    pane, qtbot, tmp_path
):
    pane.load(StubView(build_notes(tmp_path, "a.docx")))
    pane.show()
    qtbot.waitExposed(pane)

    assert pane.highlight_matches("NOTHING MATCHES THIS") == 0
    assert pane._stack.currentWidget() is pane._editor


def test_queue_highlight_is_debounced_and_fires_once_typing_stops(
    pane, qtbot, tmp_path
):
    pane.load(StubView(build_notes(tmp_path, "a.docx", rows=4, text="TORQUE")))
    pane.show()
    qtbot.waitExposed(pane)

    pane.queue_highlight("TOR")
    pane.queue_highlight("TORQ")
    pane.queue_highlight("TORQUE")
    assert pane._editor.extraSelections() == [], "highlighting ran before the pause"

    qtbot.waitUntil(lambda: len(pane._editor.extraSelections()) == 4, timeout=2000)


def test_a_pending_query_is_reapplied_when_the_pane_renders(pane, qtbot, tmp_path):
    pane.highlight_matches("TORQUE")  # Empty pane; remembered, not applied
    pane.load(StubView(build_notes(tmp_path, "a.docx", rows=3, text="TORQUE")))
    pane.show()
    qtbot.waitExposed(pane)
    assert len(pane._editor.extraSelections()) == 3


# ------------------------------------------------------------------ clicks


def test_a_click_outside_every_table_emits_empty_space_clicked(
    pane, qtbot, tmp_path
):
    pane.load(StubView(build_notes(tmp_path, "a.docx")))
    pane.resize(900, 700)
    pane.show()
    qtbot.waitExposed(pane)

    below_the_table = QPoint(
        pane._editor.viewport().width() - 2,
        pane._editor.viewport().height() - 2,
    )
    assert pane._table_at(below_the_table) is None
    with qtbot.waitSignal(pane.empty_space_clicked, timeout=1000):
        pane._on_viewport_press(below_the_table)


def test_a_click_inside_a_table_does_not_emit(pane, qtbot, tmp_path):
    pane.load(StubView(build_notes(tmp_path, "a.docx")))
    pane.show()
    qtbot.waitExposed(pane)

    inside = QPoint(6, 6)
    if pane._table_at(inside) is None:
        pytest.skip("layout did not place a table at the probe point")

    fired = []
    pane.empty_space_clicked.connect(lambda: fired.append(True))
    assert pane._on_viewport_press(inside) is False
    assert fired == []


def test_the_context_menu_predicate_is_the_same_one_used_for_clicks(
    pane, qtbot, tmp_path
):
    pane.load(StubView(build_notes(tmp_path, "a.docx")))
    pane.resize(900, 700)
    pane.show()
    qtbot.waitExposed(pane)

    outside = QPoint(
        pane._editor.viewport().width() - 2,
        pane._editor.viewport().height() - 2,
    )
    assert pane._table_at(outside) is None
    pane._on_context_menu(outside)  # no table -> no menu, and no raise


def test_a_document_with_no_declared_widths_wraps_to_the_widget(pane, qtbot, tmp_path):
    from PyQt6.QtWidgets import QTextEdit

    document = docx.Document()
    table = document.add_table(1, 3)
    tbl = table._tbl
    tbl.remove(tbl.tblGrid)
    path = tmp_path / "nogrid.docx"
    document.save(str(path))

    pane.load(StubView(path))
    pane.show()
    qtbot.waitExposed(pane)
    assert pane._editor.lineWrapMode() == QTextEdit.LineWrapMode.WidgetWidth


def test_a_fixed_width_document_gets_its_natural_text_width(pane, qtbot, tmp_path):
    from docx.oxml.ns import qn

    document = docx.Document()
    table = document.add_table(1, 2)
    grid = table._tbl.tblGrid
    for element in list(grid):
        grid.remove(element)
    for twips in (450, 9000):
        column = grid.makeelement(qn("w:gridCol"), {})
        column.set(qn("w:w"), str(twips))
        grid.append(column)
    path = tmp_path / "wide.docx"
    document.save(str(path))

    pane.load(StubView(path))
    pane.show()
    qtbot.waitExposed(pane)

    from PyQt6.QtWidgets import QTextEdit

    assert pane._editor.lineWrapMode() == QTextEdit.LineWrapMode.FixedPixelWidth
    assert pane._editor.lineWrapColumnOrWidth() == 630
