"""Patch 08 §3 -- the build-notes renderer, against synthetic .docx fixtures."""

import io
import zipfile

import docx
import pytest
from docx.oxml.ns import qn
from docx.shared import Emu
from PyQt6.QtGui import QImage, QTextDocument, QTextLength

from cockpit.services.notes_renderer import (
    MAX_DOCUMENT_MEDIA_PX,
    NotesRenderFailure,
    RenderPalette,
    RenderedNotes,
    render_build_notes,
    resolve_cell_borders,
    resolve_column_constraints,
)

PALETTE = RenderPalette(
    page_background_rgb="#FFFFFF",
    default_text_rgb="#000000",
    placeholder_border_rgb="#CC4444",
    placeholder_text_rgb="#996666",
)

EMU_PER_PX = 9525
TWIPS_PER_PX = 15


@pytest.fixture(autouse=True)
def _qt_app(qapp):
    """Every QImage/QTextDocument here needs a live QGuiApplication."""
    return qapp


# ------------------------------------------------------------------ helpers


def save(document, tmp_path, name):
    path = tmp_path / name
    document.save(str(path))
    return path


def png_bytes(width, height, colour=0xFF3366CC):
    image = QImage(width, height, QImage.Format.Format_ARGB32)
    image.fill(colour)
    buffer = io.BytesIO()
    from PyQt6.QtCore import QBuffer, QByteArray, QIODevice

    payload = QByteArray()
    qbuffer = QBuffer(payload)
    qbuffer.open(QIODevice.OpenModeFlag.WriteOnly)
    image.save(qbuffer, "PNG")
    qbuffer.close()
    buffer.write(bytes(payload))
    return buffer.getvalue()


def set_grid(table, twips_per_column):
    """Rewrites w:tblGrid; a None entry omits w:w for that column."""
    grid = table._tbl.tblGrid
    for element in list(grid):
        grid.remove(element)
    for twips in twips_per_column:
        col = grid.makeelement(qn("w:gridCol"), {})
        if twips is not None:
            col.set(qn("w:w"), str(twips))
        grid.append(col)


def drop_grid(table):
    tbl = table._tbl
    grid = tbl.tblGrid
    if grid is not None:
        tbl.remove(grid)


def set_table_width(table, twips):
    properties = table._tbl.tblPr
    existing = properties.find(qn("w:tblW"))
    if existing is not None:
        properties.remove(existing)
    element = properties.makeelement(qn("w:tblW"), {})
    element.set(qn("w:type"), "dxa")
    element.set(qn("w:w"), str(twips))
    properties.append(element)


def tc_at(table, row, index):
    """Raw w:tc by position; Table.cell() cannot address a malformed merge."""
    return table._tbl.tr_lst[row].tc_lst[index]


def set_vmerge(cell, value):
    properties = _tc_of(cell).get_or_add_tcPr()
    element = properties.makeelement(qn("w:vMerge"), {})
    element.set(qn("w:val"), value)
    properties.append(element)


def _tc_of(cell):
    return cell if cell.__class__.__name__ == "CT_Tc" else cell._tc


def set_grid_span(cell, span):
    properties = _tc_of(cell).get_or_add_tcPr()
    element = properties.makeelement(qn("w:gridSpan"), {})
    element.set(qn("w:val"), str(span))
    properties.append(element)


def set_borders(element_parent, tag, edges):
    """edges: {'top': 'single'|'nil', ...}"""
    borders = element_parent.makeelement(qn(tag), {})
    for edge, value in edges.items():
        edge_element = borders.makeelement(qn(f"w:{edge}"), {})
        edge_element.set(qn("w:val"), value)
        borders.append(edge_element)
    element_parent.append(borders)


def add_picture(cell, blob, display_px=None):
    """Inserts an inline picture, optionally overriding wp:extent."""
    run = cell.paragraphs[0].add_run()
    run.add_picture(io.BytesIO(blob))
    if display_px is not None:
        width, height = display_px
        drawing = run._r.find(qn("w:drawing"))
        for extent in drawing.iter():
            if extent.tag.endswith("}extent") or extent.tag.endswith("}ext"):
                extent.set("cx", str(width * EMU_PER_PX))
                extent.set("cy", str(height * EMU_PER_PX))
    return run


def replace_image_blob(path, blob, content_type=None):
    """Rewrites the single media part of a saved .docx in place."""
    source = zipfile.ZipFile(path)
    entries = [(item, source.read(item.filename)) for item in source.infolist()]
    source.close()

    rewritten = []
    for item, data in entries:
        if item.filename.startswith("word/media/"):
            data = blob
        if content_type and item.filename == "[Content_Types].xml":
            data = data.replace(b'ContentType="image/png"',
                                f'ContentType="{content_type}"'.encode())
        rewritten.append((item, data))

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as target:
        for item, data in rewritten:
            target.writestr(item, data)
    return path


def only_table(rendered: RenderedNotes):
    cursor = rendered.document.rootFrame().begin()
    block = rendered.document.begin()
    while block.isValid():
        from PyQt6.QtGui import QTextCursor

        probe = QTextCursor(block)
        table = probe.currentTable()
        if table is not None:
            return table
        block = block.next()
    raise AssertionError("no table in rendered document")


def tables_of(rendered: RenderedNotes):
    from PyQt6.QtGui import QTextCursor

    found = []
    block = rendered.document.begin()
    while block.isValid():
        table = QTextCursor(block).currentTable()
        if table is not None and all(table != seen for seen in found):
            found.append(table)
        block = block.next()
    return found


# ------------------------------------------------------------------ failures


def test_document_missing_returns_rather_than_raises(tmp_path):
    outcome = render_build_notes(tmp_path / "absent.docx", PALETTE)
    assert not outcome.is_ok()
    assert outcome.err.reason == "DocumentMissing"
    assert "absent.docx" in outcome.err.message


def test_document_unreadable_returns_the_underlying_reason(tmp_path):
    path = tmp_path / "corrupt.docx"
    path.write_bytes(b"PK\x03\x04 not really a docx")
    outcome = render_build_notes(path, PALETTE)
    assert not outcome.is_ok()
    assert outcome.err.reason == "DocumentUnreadable"
    assert outcome.err.message


def test_no_tables_found(tmp_path):
    document = docx.Document()
    document.add_paragraph("Prose only, no tables.")
    outcome = render_build_notes(save(document, tmp_path, "prose.docx"), PALETTE)
    assert not outcome.is_ok()
    assert outcome.err.reason == "NoTablesFound"


def test_a_directory_in_place_of_the_document_is_missing_not_a_crash(tmp_path):
    directory = tmp_path / "notes.docx"
    directory.mkdir()
    outcome = render_build_notes(directory, PALETTE)
    assert not outcome.is_ok()
    assert outcome.err.reason == "DocumentMissing"


# ------------------------------------------------------------------ structure


def test_plain_table_renders_rows_columns_and_text(tmp_path):
    document = docx.Document()
    table = document.add_table(2, 3)
    table.cell(0, 0).text = "Step"
    table.cell(1, 2).text = "Torque to 4 in-lb"

    outcome = render_build_notes(save(document, tmp_path, "plain.docx"), PALETTE)
    assert outcome.is_ok()
    rendered = only_table(outcome.ok)
    assert (rendered.rows(), rendered.columns()) == (2, 3)
    assert "Step" in outcome.ok.document.toPlainText()
    assert "Torque to 4 in-lb" in outcome.ok.document.toPlainText()


def test_every_source_table_renders_in_document_order(tmp_path):
    document = docx.Document()
    first = document.add_table(1, 1)
    first.cell(0, 0).text = "FIRST"
    document.add_paragraph("")
    second = document.add_table(1, 1)
    second.cell(0, 0).text = "SECOND"

    outcome = render_build_notes(save(document, tmp_path, "two.docx"), PALETTE)
    assert outcome.is_ok()
    assert len(tables_of(outcome.ok)) == 2
    text = outcome.ok.document.toPlainText()
    assert text.index("FIRST") < text.index("SECOND")


# ------------------------------------------------------------------ geometry


def test_fully_specified_grid_gives_fixed_widths_and_natural_width(tmp_path):
    document = docx.Document()
    table = document.add_table(1, 2)
    set_grid(table, [450, 9000])

    outcome = render_build_notes(save(document, tmp_path, "grid.docx"), PALETTE)
    assert outcome.is_ok()
    constraints = only_table(outcome.ok).format().columnWidthConstraints()
    assert [c.type() for c in constraints] == [QTextLength.Type.FixedLength] * 2
    assert [round(c.rawValue()) for c in constraints] == [30, 600]
    assert outcome.ok.natural_width == 630


def test_grid_is_honoured_even_when_the_table_declares_autofit(tmp_path):
    """Word's autofit still lays out from tblGrid; equal columns are wrong."""
    document = docx.Document()
    table = document.add_table(1, 3)
    table.autofit = True
    set_grid(table, [709, 2007, 3706])

    outcome = render_build_notes(save(document, tmp_path, "autofit.docx"), PALETTE)
    assert outcome.is_ok()
    constraints = only_table(outcome.ok).format().columnWidthConstraints()
    assert [c.type() for c in constraints] == [QTextLength.Type.FixedLength] * 3
    assert outcome.ok.natural_width > 0


def test_partial_grid_splits_the_declared_table_width_among_the_unspecified(tmp_path):
    document = docx.Document()
    table = document.add_table(1, 3)
    set_grid(table, [1500, None, None])
    set_table_width(table, 4500)

    constraints, natural = resolve_column_constraints(table, 3)
    assert [c.type() for c in constraints] == [QTextLength.Type.FixedLength] * 3
    assert round(constraints[0].rawValue()) == 100
    assert round(constraints[1].rawValue()) == 100
    assert round(constraints[2].rawValue()) == 100
    assert round(natural) == 300


def test_partial_grid_with_no_declared_table_width_uses_the_mean(tmp_path):
    document = docx.Document()
    table = document.add_table(1, 3)
    set_grid(table, [1500, 3000, None])
    set_table_width(table, 0)

    constraints, natural = resolve_column_constraints(table, 3)
    assert round(constraints[2].rawValue()) == 150  # mean of 100 and 200
    assert round(natural) == 450


def test_absent_grid_falls_back_to_even_percentages_and_zero_natural_width(tmp_path):
    document = docx.Document()
    table = document.add_table(1, 4)
    drop_grid(table)

    outcome = render_build_notes(save(document, tmp_path, "nogrid.docx"), PALETTE)
    assert outcome.is_ok()
    constraints = only_table(outcome.ok).format().columnWidthConstraints()
    assert [c.type() for c in constraints] == [QTextLength.Type.PercentageLength] * 4
    assert [c.rawValue() for c in constraints] == [25.0] * 4
    assert outcome.ok.natural_width == 0


def test_natural_width_counts_only_fixed_width_tables(tmp_path):
    document = docx.Document()
    fixed = document.add_table(1, 2)
    set_grid(fixed, [1500, 1500])
    document.add_paragraph("")
    floating = document.add_table(1, 6)
    drop_grid(floating)

    outcome = render_build_notes(save(document, tmp_path, "mixed.docx"), PALETTE)
    assert outcome.is_ok()
    assert outcome.ok.natural_width == 200


# ------------------------------------------------------------------ merges


def test_grid_span_merges_horizontally(tmp_path):
    document = docx.Document()
    table = document.add_table(2, 3)
    set_grid_span(tc_at(table, 0, 0), 2)

    outcome = render_build_notes(save(document, tmp_path, "hspan.docx"), PALETTE)
    assert outcome.is_ok()
    rendered = only_table(outcome.ok)
    assert rendered.cellAt(0, 0).columnSpan() == 2
    assert outcome.ok.anomalies == []


def test_vmerge_run_of_three_rows_merges_vertically(tmp_path):
    document = docx.Document()
    table = document.add_table(4, 2)
    set_vmerge(table.cell(0, 0), "restart")
    set_vmerge(table.cell(1, 0), "continue")
    set_vmerge(table.cell(2, 0), "continue")
    table.cell(0, 0).text = "SPANNING"

    outcome = render_build_notes(save(document, tmp_path, "vmerge.docx"), PALETTE)
    assert outcome.is_ok()
    rendered = only_table(outcome.ok)
    assert rendered.cellAt(0, 0).rowSpan() == 3
    assert outcome.ok.anomalies == []


def test_vmerge_is_applied_after_the_walk_and_does_not_scramble_later_columns(tmp_path):
    """The columns right of a vertical merge must keep their own text."""
    document = docx.Document()
    table = document.add_table(3, 3)
    for row in range(3):
        for column in range(3):
            table.cell(row, column).text = f"r{row}c{column}"
    set_vmerge(table.cell(0, 1), "restart")
    set_vmerge(table.cell(1, 1), "continue")
    set_vmerge(table.cell(2, 1), "continue")

    outcome = render_build_notes(save(document, tmp_path, "vscramble.docx"), PALETTE)
    assert outcome.is_ok()
    rendered = only_table(outcome.ok)
    for row in range(3):
        cell = rendered.cellAt(row, 2)
        text = cell.firstCursorPosition().block().text()
        assert text == f"r{row}c2", f"column 2 scrambled at row {row}: {text!r}"


def test_a_cell_merged_both_ways_keeps_both_spans(tmp_path):
    document = docx.Document()
    table = document.add_table(3, 3)
    anchor = table.cell(0, 0)
    set_grid_span(anchor, 2)
    set_vmerge(anchor, "restart")
    set_vmerge(tc_at(table, 1, 0), "continue")

    outcome = render_build_notes(save(document, tmp_path, "both.docx"), PALETTE)
    assert outcome.is_ok()
    rendered = only_table(outcome.ok)
    assert rendered.cellAt(0, 0).rowSpan() == 2
    assert rendered.cellAt(0, 0).columnSpan() == 2
    assert [a.reason for a in outcome.ok.anomalies] == []


def test_orphaned_vmerge_continue_is_an_anomaly_not_an_exception(tmp_path):
    document = docx.Document()
    table = document.add_table(2, 2)
    table.cell(1, 1).text = "STILL READABLE"
    set_vmerge(tc_at(table, 0, 0), "continue")  # no restart above it

    outcome = render_build_notes(save(document, tmp_path, "orphan.docx"), PALETTE)
    assert outcome.is_ok()
    assert [a.reason for a in outcome.ok.anomalies] == ["MergeConflict"]
    assert "STILL READABLE" in outcome.ok.document.toPlainText()


def test_overlapping_merges_yield_a_conflict_and_a_readable_table(tmp_path):
    document = docx.Document()
    table = document.add_table(3, 3)
    for row in range(3):
        for column in range(3):
            table.cell(row, column).text = f"r{row}c{column}"
    # Row 1 column 0 spans two columns and opens a vertical run; column 1
    # opens its own run inside the region the span already claims.
    set_grid_span(tc_at(table, 1, 0), 2)
    set_vmerge(tc_at(table, 1, 0), "restart")
    set_vmerge(tc_at(table, 2, 0), "continue")
    set_vmerge(tc_at(table, 1, 1), "restart")
    set_vmerge(tc_at(table, 2, 1), "continue")

    outcome = render_build_notes(save(document, tmp_path, "conflict.docx"), PALETTE)
    assert outcome.is_ok()
    assert any(a.reason == "MergeConflict" for a in outcome.ok.anomalies)
    assert "r0c0" in outcome.ok.document.toPlainText()


# ------------------------------------------------------------------ borders


def test_cell_borders_override_table_borders_per_edge(tmp_path):
    document = docx.Document()
    table = document.add_table(1, 1)
    set_borders(table._tbl.tblPr, "w:tblBorders",
                {"top": "single", "bottom": "single", "left": "single", "right": "single"})
    cell = table.cell(0, 0)
    set_borders(cell._tc.get_or_add_tcPr(), "w:tcBorders", {"top": "double"})

    resolved = resolve_cell_borders(cell, table)
    assert resolved == {"top": True, "bottom": True, "left": True, "right": True}


def test_a_nil_cell_edge_suppresses_rather_than_inheriting(tmp_path):
    document = docx.Document()
    table = document.add_table(1, 1)
    set_borders(table._tbl.tblPr, "w:tblBorders",
                {"top": "single", "bottom": "single", "left": "single", "right": "single"})
    cell = table.cell(0, 0)
    set_borders(cell._tc.get_or_add_tcPr(), "w:tcBorders", {"top": "nil"})

    resolved = resolve_cell_borders(cell, table)
    assert resolved["top"] is False, "nil on a cell edge must not fall through"
    assert resolved["bottom"] is True


def test_absent_borders_everywhere_means_no_border(tmp_path):
    document = docx.Document()
    table = document.add_table(1, 1)
    assert resolve_cell_borders(table.cell(0, 0), table) == {
        "top": False, "bottom": False, "left": False, "right": False
    }


# ------------------------------------------------------------------ images


def test_image_is_decoded_to_the_extent_word_declares_not_its_native_size(tmp_path):
    """The regression guard for scaled decoding (§3.4)."""
    document = docx.Document()
    table = document.add_table(1, 1)
    add_picture(table.cell(0, 0), png_bytes(800, 600), display_px=(200, 150))

    outcome = render_build_notes(save(document, tmp_path, "scaled.docx"), PALETTE)
    assert outcome.is_ok()
    assert len(outcome.ok.image_names) == 1
    name = outcome.ok.image_names[0]
    stored = outcome.ok.document.resource(
        QTextDocument.ResourceType.ImageResource, _url(name)
    )
    image = stored if isinstance(stored, QImage) else QImage(stored)
    assert (image.width(), image.height()) == (200, 150)
    assert outcome.ok.anomalies == []


def test_an_image_with_no_extent_falls_back_to_its_intrinsic_size(tmp_path):
    document = docx.Document()
    table = document.add_table(1, 1)
    run = table.cell(0, 0).paragraphs[0].add_run()
    run.add_picture(io.BytesIO(png_bytes(120, 90)))
    drawing = run._r.find(qn("w:drawing"))
    for extent in list(drawing.iter()):
        if extent.tag.endswith("}extent"):
            extent.getparent().remove(extent)

    outcome = render_build_notes(save(document, tmp_path, "noextent.docx"), PALETTE)
    assert outcome.is_ok()
    name = outcome.ok.image_names[0]
    stored = outcome.ok.document.resource(
        QTextDocument.ResourceType.ImageResource, _url(name)
    )
    image = stored if isinstance(stored, QImage) else QImage(stored)
    assert (image.width(), image.height()) == (120, 90)


def test_image_resource_names_are_ours_and_prefixed(tmp_path):
    document = docx.Document()
    table = document.add_table(1, 1)
    add_picture(table.cell(0, 0), png_bytes(40, 40))

    outcome = render_build_notes(save(document, tmp_path, "named.docx"), PALETTE)
    assert outcome.is_ok()
    assert all(name.startswith("notes-img:") for name in outcome.ok.image_names)


def test_an_undecodable_medium_is_placeheld_and_recorded(tmp_path):
    document = docx.Document()
    table = document.add_table(1, 2)
    add_picture(table.cell(0, 0), png_bytes(40, 40))
    table.cell(0, 1).text = "REST OF THE TABLE"
    path = save(document, tmp_path, "emf.docx")
    replace_image_blob(path, b"\x01\x00\x00\x00 not an image", content_type="image/x-emf")

    outcome = render_build_notes(path, PALETTE)
    assert outcome.is_ok()
    assert [a.reason for a in outcome.ok.anomalies] == ["UndecodableMedium"]
    assert "REST OF THE TABLE" in outcome.ok.document.toPlainText()


def test_a_blob_that_is_not_the_image_it_claims_is_a_degenerate_decode(tmp_path):
    document = docx.Document()
    table = document.add_table(1, 2)
    add_picture(table.cell(0, 0), png_bytes(40, 40), display_px=(40, 40))
    table.cell(0, 1).text = "REST OF THE TABLE"
    path = save(document, tmp_path, "truncated.docx")
    replace_image_blob(path, png_bytes(40, 40)[:20])  # truncated PNG, still image/png

    outcome = render_build_notes(path, PALETTE)
    assert outcome.is_ok()
    assert [a.reason for a in outcome.ok.anomalies] == ["DegenerateImage"]
    assert outcome.ok.image_names  # the placeholder is registered
    assert not any(n.startswith("notes-img:") for n in outcome.ok.image_names)
    assert "REST OF THE TABLE" in outcome.ok.document.toPlainText()


def test_media_budget_placeholds_rather_than_decoding(tmp_path, monkeypatch):
    import cockpit.services.notes_renderer as renderer

    monkeypatch.setattr(renderer, "MAX_DOCUMENT_MEDIA_PX", 100)
    document = docx.Document()
    table = document.add_table(1, 1)
    add_picture(table.cell(0, 0), png_bytes(200, 200), display_px=(200, 200))

    outcome = render_build_notes(save(document, tmp_path, "budget.docx"), PALETTE)
    assert outcome.is_ok()
    assert [a.reason for a in outcome.ok.anomalies] == ["MediaBudgetExceeded"]


# ------------------------------------------------------------------ formatting


def test_a_run_without_an_explicit_colour_takes_the_palette_default(tmp_path):
    from PyQt6.QtGui import QColor, QTextCursor

    document = docx.Document()
    table = document.add_table(1, 1)
    table.cell(0, 0).paragraphs[0].add_run("UNCOLOURED")

    palette = RenderPalette("#FFFFFF", "#123456", "#CC4444", "#996666")
    outcome = render_build_notes(save(document, tmp_path, "colour.docx"), palette)
    assert outcome.is_ok()
    cell = only_table(outcome.ok).cellAt(0, 0)
    cursor = cell.firstCursorPosition()
    cursor.movePosition(QTextCursor.MoveOperation.NextCharacter,
                        QTextCursor.MoveMode.KeepAnchor)
    assert cursor.charFormat().foreground().color() == QColor("#123456")


def test_bold_and_italic_survive(tmp_path):
    from PyQt6.QtGui import QFont, QTextCursor

    document = docx.Document()
    table = document.add_table(1, 1)
    run = table.cell(0, 0).paragraphs[0].add_run("EMPHATIC")
    run.bold = True
    run.italic = True

    outcome = render_build_notes(save(document, tmp_path, "bold.docx"), PALETTE)
    assert outcome.is_ok()
    cursor = only_table(outcome.ok).cellAt(0, 0).firstCursorPosition()
    cursor.movePosition(QTextCursor.MoveOperation.NextCharacter,
                        QTextCursor.MoveMode.KeepAnchor)
    assert cursor.charFormat().fontWeight() == QFont.Weight.Bold
    assert cursor.charFormat().fontItalic()


def test_tabs_and_breaks_render_as_text_not_element_repr(tmp_path):
    document = docx.Document()
    table = document.add_table(1, 1)
    run = table.cell(0, 0).paragraphs[0].add_run("BEFORE")
    run.add_tab()
    run.add_text("AFTER")

    outcome = render_build_notes(save(document, tmp_path, "tabs.docx"), PALETTE)
    assert outcome.is_ok()
    text = outcome.ok.document.toPlainText()
    assert "BEFORE\tAFTER" in text
    assert "Element" not in text


def _url(name):
    from PyQt6.QtCore import QUrl

    return QUrl(name)
